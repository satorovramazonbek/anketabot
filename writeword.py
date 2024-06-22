import docx
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENTATION
from database import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
doc = docx.Document()
def setting():
    section = doc.sections[-1]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.buttom_margin = Inches(0.5)
async def write(user_id, dict_data, a):
    setting()
    title = doc.add_paragraph('\t\t\t\t\t\tMa\'lumotnoma')
    title.runs[0].font.size = Pt(18)
    title.style = 'Title'
    text_paragraph = doc.add_paragraph('\t\t\t\t\t\t\t\t\t\t\t')
    text_paragraph.add_run().add_picture(f"{user_id}.jpg", width=Inches(1.5), height=Inches(2.0))

    # fm = doc.add_paragraph().add_run(f"\t\t\t\t{dict_data[user_id]['fio']}")
    fm = doc.add_paragraph().add_run(f"\t\t\t\t{await select_user(user_id,'fio')}")
    fm.font.size = Pt(14)
    fm.bold = True

    # fm = doc.add_paragraph().add_run(f"{dict_data[user_id]['study']}")
    fm = doc.add_paragraph().add_run(f"{await select_user(user_id,'study')}")
    fm.font.size = Pt(12)
    fm.bold = True


    fm2 = doc.add_paragraph().add_run('\nTug‘ilgan yili:\t\t\t\t\t\tTug‘ilgan joyi:')
    # fm3 = doc.add_paragraph().add_run(f"{dict_data[user_id]['birthday']}:\t\t\t\t\t\t\t{dict_data[user_id]['country']}")
    fm3 = doc.add_paragraph().add_run(f"{await select_user(user_id,'birthday')}:\t\t\t\t\t\t\t{await select_user(user_id,'country')}")
    fm4 = doc.add_paragraph().add_run('\nMillati:\t\t\t\t\t\t\tPartiyaviyligi:')
    # fm5 = doc.add_paragraph().add_run(f"{dict_data[user_id]['millati']}:\t\t\t\t\t\t\t{dict_data[user_id]['partiyaviylik']}")
    fm5 = doc.add_paragraph().add_run(f"{await select_user(user_id,'millati')}:\t\t\t\t\t\t\t{await select_user(user_id,'partiyaviylik')}")
    # fm6 = doc.add_paragraph().add_run(f"\nMa’lumoti:  {dict_data[user_id]['malumoti']}")  #data_dict dan ma'lumotni olib qoyiladi
    fm6 = doc.add_paragraph().add_run(f"\nMa’lumoti:  {await select_user(user_id,'malumoti')}")  #data_dict dan ma'lumotni olib qoyiladi
    # fm7 = doc.add_paragraph().add_run(f"Tamomlagan:  {dict_data[user_id]['tamomlagan']}")
    fm7 = doc.add_paragraph().add_run(f"Tamomlagan:  {await select_user(user_id,'tamomlagan')}")
    # fm8 = doc.add_paragraph().add_run(f"Ma’lumoti bo‘yicha mutaxassisligi:    {dict_data[user_id]['mutaxassislik']} ")
    fm8 = doc.add_paragraph().add_run(f"Ma’lumoti bo‘yicha mutaxassisligi:    {await select_user(user_id,'mutaxassislik')} ")
    # fm9 = doc.add_paragraph().add_run(f"Qaysi chet tillarini biladi:   {dict_data[user_id]['language']}")
    fm9 = doc.add_paragraph().add_run(f"Qaysi chet tillarini biladi:   {await select_user(user_id,'language')}")
    # fm10 = doc.add_paragraph().add_run(f"Davlat mukofotlari bilan taqdirlanganmi (qanaqa):  {dict_data[user_id]['mukofot']}") # yoki pastidan qushsa buladi
    fm10 = doc.add_paragraph().add_run(f"Davlat mukofotlari bilan taqdirlanganmi (qanaqa):  {await select_user(user_id,'mukofot')}") # yoki pastidan qushsa buladi
#     fm11 = doc.add_paragraph().add_run(f"""Xalq  deputatlari, respublika, viloyat, shahar va tuman Kengashi deputatimi yoki boshqa
# saylanadigan organlarning a’zosimi (to‘liq ko‘rsatilishi lozim):	 {dict_data[user_id]['saylov']} """)
    fm11 = doc.add_paragraph().add_run(f"""Xalq  deputatlari, respublika, viloyat, shahar va tuman Kengashi deputatimi yoki boshqa
saylanadigan organlarning a’zosimi (to‘liq ko‘rsatilishi lozim):	 {await select_user(user_id,'saylov')} """)
    fm12 = doc.add_paragraph().add_run('\t\t\t\t\t\tMЕHNAT FAOLIYATI ')
    # fm13 = doc.add_paragraph().add_run(f"{dict_data[user_id]['mehnat_faoliyati']}")
    fm13 = doc.add_paragraph().add_run(f"{await select_user(user_id,'mehnat_faoliyati')}")
    fm2.font.size = Pt(12)
    fm3.font.size = Pt(12)
    fm4.font.size = Pt(12)
    fm5.font.size = Pt(12)
    fm6.font.size = Pt(12)
    fm7.font.size = Pt(12)
    fm8.font.size = Pt(12)
    fm9.font.size = Pt(12)
    fm10.font.size = Pt(12)
    fm11.font.size = Pt(12)
    fm12.font.size = Pt(14)
    fm13.font.size = Pt(12)
    fm2.bold = True
    fm4.bold = True
    fm5.bold = True
    fm6.bold = True
    fm7.bold = True
    fm8.bold = True
    fm9.bold = True
    fm10.bold = True
    fm11.bold = True
    fm12.bold = True

    doc.add_page_break()
    # doc.add_paragraph().add_run(f"\t\t\t{dict_data['fio']}ning yaqin qarindoshlari haqida\n\t\t\t\t\t\t\tMA’LUMOT").bold = True
    doc.add_paragraph().add_run(f"\t\t\t{await select_user(user_id,'fio')}ning yaqin qarindoshlari haqida\n\t\t\t\t\t\t\tMA’LUMOT").bold = True
    table = doc.add_table(rows=1, cols=5)

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qarindoshligi'
    hdr_cells[1].text = 'Familiyasi, ismi va otasining ismi'
    hdr_cells[2].text = 'Tug‘ilgan yili va joyi'
    hdr_cells[3].text = 'Ish joyi va lavozimi'
    hdr_cells[4].text = 'Turar joyi'
    for id in range(1, a+1):
        row_cells = table.add_row().cells
        nds = await select_user(user_id, 'qarindoshlari')
        parent_list = nds.split(',')
        # row_cells[0].text = dict_data[f"qarindoshlik_turi{id}"]  # Преобразование id в строку
        row_cells[0].text = await sele_parent(user_id, 'qarindoshligi', parent_list[a-1])
        row_cells[1].text = await sele_parent(user_id, "fio", parent_list[a-1])
        row_cells[2].text = await sele_parent(user_id, 'birthday', parent_list[a-1])
        row_cells[3].text =await sele_parent(user_id, 'work', parent_list[a-1])
        row_cells[4].text = await sele_parent(user_id, 'location', parent_list[a-1])



    for cell in hdr_cells:
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.size = Pt(12)
    doc.save(f"{user_id}.docx")
